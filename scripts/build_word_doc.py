#!/usr/bin/env python
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a Word document from the course paper LaTeX source.")
    parser.add_argument("--tex", default="paper/main.tex")
    parser.add_argument("--output", default="paper/main.docx")
    return parser.parse_args()


def set_run_font(run, *, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.05)
    section.right_margin = Inches(1.05)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Title"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    doc.styles["Normal"].font.size = Pt(11)


def strip_outer_braces(text: str) -> str:
    text = text.strip()
    if text.startswith("{") and text.endswith("}"):
        return text[1:-1].strip()
    return text


def replace_simple_commands(text: str) -> str:
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\(?:texttt|textbf|emph|heiti)\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\(?:mathrm|hat)\{([^{}]*)\}", r"\1", text)
    return text


def clean_tex(text: str, refs: dict[str, str] | None = None, cites: dict[str, int] | None = None) -> str:
    refs = refs or {}
    cites = cites or {}

    def cite_repl(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        values = [str(cites.get(key, key)) for key in keys]
        return "[" + ",".join(values) + "]"

    def ref_repl(match: re.Match[str]) -> str:
        return refs.get(match.group(1), "?")

    text = text.strip()
    text = text.replace("~", " ")
    text = text.replace("\\quad", " ")
    replacements = {
        "\\alpha": "α",
        "\\theta": "θ",
        "\\Delta": "Δ",
        "\\sigma": "σ",
        "\\min": "min",
        "\\log": "log",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"\\cite\{([^}]*)\}", cite_repl, text)
    text = re.sub(r"\\ref\{([^}]*)\}", ref_repl, text)
    text = re.sub(r"\\url\{([^}]*)\}", r"\1", text)
    text = replace_simple_commands(text)
    text = text.replace("\\%", "%").replace("\\_", "_").replace("\\#", "#").replace("\\&", "&")
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = re.sub(r"\$(.*?)\$", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def collect_references(tex: str) -> dict[str, int]:
    refs: dict[str, int] = {}
    for index, key in enumerate(re.findall(r"\\bibitem\{([^}]*)\}", tex), start=1):
        refs[key] = index
    return refs


def collect_labels(tex: str) -> dict[str, str]:
    labels: dict[str, str] = {}
    counters = {"table": 0, "figure": 0, "equation": 0}
    block_pattern = re.compile(r"\\begin\{(table|figure|equation)\}.*?\\end\{\1\}", re.S)
    for match in block_pattern.finditer(tex):
        block_type = match.group(1)
        block = match.group(0)
        counters[block_type] += 1
        for label in re.findall(r"\\label\{([^}]*)\}", block):
            labels[label] = str(counters[block_type])
    return labels


def add_centered(doc: Document, text: str, *, size: int = 11, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_paragraph(doc: Document, text: str, refs: dict[str, str], cites: dict[str, int]) -> None:
    text = clean_tex(text, refs, cites)
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_list_item(doc: Document, text: str, refs: dict[str, str], cites: dict[str, int], *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(clean_tex(text, refs, cites))
    set_run_font(run, size=11)


def collect_toc_entries(body: str, refs: dict[str, str], cites: dict[str, int]) -> list[tuple[int, str, str]]:
    entries: list[tuple[int, str, str]] = []
    section_no = 0
    subsection_no = 0
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if line.startswith("\\section"):
            section_no += 1
            subsection_no = 0
            title_match = re.search(r"\\section\{(.+?)\}", line)
            if title_match:
                entries.append((1, str(section_no), clean_tex(title_match.group(1), refs, cites)))
        elif line.startswith("\\subsection"):
            subsection_no += 1
            title_match = re.search(r"\\subsection\{(.+?)\}", line)
            if title_match:
                entries.append((2, f"{section_no}.{subsection_no}", clean_tex(title_match.group(1), refs, cites)))
    return entries


def add_table_of_contents(doc: Document, entries: list[tuple[int, str, str]]) -> None:
    add_centered(doc, "目录", size=14, bold=True)
    doc.add_paragraph()
    for level, number, title in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(0 if level == 1 else 18)
        paragraph.paragraph_format.line_spacing = 1.2
        run = paragraph.add_run(f"{number} {title}")
        set_run_font(run, size=11 if level == 1 else 10, bold=(level == 1))


def resolve_image_path(raw_path: str) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path
    return (ROOT / "paper" / path).resolve()


def add_image_to_paragraph(paragraph, image_path: Path, width_inches: float) -> None:
    if image_path.exists():
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))


def parse_caption(block: str) -> str:
    match = re.search(r"\\caption\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}", block, re.S)
    return clean_tex(match.group(1), {}, {}) if match else ""


def parse_label(block: str) -> str | None:
    match = re.search(r"\\label\{([^}]*)\}", block)
    return match.group(1) if match else None


def add_figure(doc: Document, block: str, refs: dict[str, str], cites: dict[str, int]) -> None:
    figure_no = refs.get(parse_label(block) or "", "")
    sub_blocks = re.findall(r"\\begin\{subfigure\}.*?\\end\{subfigure\}", block, flags=re.S)
    main_caption = parse_caption(re.sub(r"\\begin\{subfigure\}.*?\\end\{subfigure\}", "", block, flags=re.S))

    if len(sub_blocks) == 2:
        table = doc.add_table(rows=2, cols=2)
        table.autofit = True
        for col, sub_block in enumerate(sub_blocks):
            image_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}", sub_block)
            if image_match:
                paragraph = table.cell(0, col).paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_image_to_paragraph(paragraph, resolve_image_path(image_match.group(1)), 2.75)
            caption = parse_caption(sub_block)
            cap_para = table.cell(1, col).paragraphs[0]
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_para.add_run(caption)
            set_run_font(run, size=9)
    else:
        for raw_path in re.findall(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}", block):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_image_to_paragraph(paragraph, resolve_image_path(raw_path), 5.7)

    if main_caption:
        add_centered(doc, f"图 {figure_no}：{clean_tex(main_caption, refs, cites)}", size=10, bold=True)


def split_latex_row(row: str) -> list[str]:
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in row:
        if char == "\\" and not escaped:
            escaped = True
            current.append(char)
            continue
        if char == "&" and not escaped:
            cells.append("".join(current))
            current = []
        else:
            current.append(char)
        escaped = False
    cells.append("".join(current))
    return cells


def add_table(doc: Document, block: str, refs: dict[str, str], cites: dict[str, int]) -> None:
    table_no = refs.get(parse_label(block) or "", "")
    caption = parse_caption(block)
    if caption:
        add_centered(doc, f"表 {table_no}：{clean_tex(caption, refs, cites)}", size=10, bold=True)

    tabular_match = re.search(r"\\begin\{tabular\}.*?\n(.*?)\\end\{tabular\}", block, re.S)
    if not tabular_match:
        return
    tabular = tabular_match.group(1)
    tabular = re.sub(r"\\(?:toprule|midrule|bottomrule)", "", tabular)
    raw_rows = [row.strip() for row in re.split(r"\\\\", tabular) if row.strip()]
    rows = []
    for raw_row in raw_rows:
        if raw_row.startswith("\\") or not raw_row:
            continue
        cells = [clean_tex(cell, refs, cites) for cell in split_latex_row(raw_row)]
        if len(cells) > 1:
            rows.append(cells)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    word_table = doc.add_table(rows=len(rows), cols=column_count)
    word_table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(column_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            paragraph = word_table.cell(r_idx, c_idx).paragraphs[0]
            run = paragraph.add_run(cell_text)
            set_run_font(run, size=8 if column_count >= 4 else 9, bold=(r_idx == 0))


def clean_equation(block: str) -> str:
    equation = re.sub(r"\\label\{[^}]*\}", "", block)
    equation = re.sub(r"\\begin\{equation\}|\\end\{equation\}", "", equation)
    equation = re.sub(r"\\begin\{aligned\}|\\end\{aligned\}", "", equation)
    equation = equation.replace("\\\\", "\n")
    equation = equation.replace("\\quad", " ")
    equation = equation.replace("&", "")
    equation = re.sub(r"\n\s*\n", "\n", equation)
    return equation.strip()


def add_equation(doc: Document, block: str, refs: dict[str, str]) -> None:
    label = parse_label(block)
    number = refs.get(label or "", "")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{clean_equation(block)}    ({number})")
    set_run_font(run, size=10)
    run.font.name = "Cambria Math"


def add_bibliography(doc: Document, tex: str, refs: dict[str, str], cites: dict[str, int]) -> None:
    doc.add_heading("参考文献", level=1)
    bib_match = re.search(r"\\begin\{thebibliography\}.*?\\end\{thebibliography\}", tex, re.S)
    if not bib_match:
        return
    content = bib_match.group(0)
    items = re.split(r"\\bibitem\{([^}]*)\}", content)[1:]
    for idx in range(0, len(items), 2):
        key = items[idx]
        text = items[idx + 1]
        number = cites.get(key, idx // 2 + 1)
        text = re.sub(r"\\end\{thebibliography\}", "", text)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[{number}] {clean_tex(text, refs, cites)}")
        set_run_font(run, size=10)


def iter_body_blocks(body: str):
    lines = body.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("\\begin{table}") or stripped.startswith("\\begin{figure}") or stripped.startswith("\\begin{equation}"):
            env = re.match(r"\\begin\{([^}]*)\}", stripped).group(1)  # type: ignore[union-attr]
            block_lines = [line]
            index += 1
            while index < len(lines) and f"\\end{{{env}}}" not in lines[index]:
                block_lines.append(lines[index])
                index += 1
            if index < len(lines):
                block_lines.append(lines[index])
            yield env, "\n".join(block_lines)
        else:
            yield "line", line
        index += 1


def build_docx(tex_path: Path, output_path: Path) -> None:
    tex = tex_path.read_text(encoding="utf-8")
    refs = collect_labels(tex)
    cites = collect_references(tex)
    doc = Document()
    configure_document(doc)

    title = clean_tex(strip_outer_braces(re.search(r"\\title\{(.+?)\}", tex, re.S).group(1)), refs, cites)  # type: ignore[union-attr]
    author = clean_tex(strip_outer_braces(re.search(r"\\author\{(.+?)\}", tex, re.S).group(1)), refs, cites)  # type: ignore[union-attr]
    add_centered(doc, title, size=18, bold=True)
    add_centered(doc, author, size=12)
    add_centered(doc, "2026 年 6 月 15 日", size=12)
    doc.add_paragraph()

    abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", tex, re.S)
    if abstract:
        add_centered(doc, "摘要", size=12, bold=True)
        for paragraph in re.split(r"\n\s*\n", abstract.group(1).strip()):
            add_paragraph(doc, paragraph, refs, cites)

    body_match = re.search(r"\\section\{引言\}(.*?)\\begin\{thebibliography\}", tex, re.S)
    if not body_match:
        raise RuntimeError("Cannot find paper body in LaTeX source.")
    body = "\\section{引言}\n" + body_match.group(1)

    doc.add_page_break()
    add_table_of_contents(doc, collect_toc_entries(body, refs, cites))
    doc.add_page_break()

    section_no = 0
    subsection_no = 0
    paragraph_lines: list[str] = []
    list_mode: str | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_paragraph(doc, " ".join(paragraph_lines), refs, cites)
            paragraph_lines = []

    for kind, payload in iter_body_blocks(body):
        if kind != "line":
            flush_paragraph()
            if kind == "table":
                add_table(doc, payload, refs, cites)
            elif kind == "figure":
                add_figure(doc, payload, refs, cites)
            elif kind == "equation":
                add_equation(doc, payload, refs)
            continue
        stripped = payload.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith("\\section"):
            flush_paragraph()
            section_no += 1
            subsection_no = 0
            title_match = re.search(r"\\section\{(.+?)\}", stripped)
            doc.add_heading(f"{section_no} {clean_tex(title_match.group(1), refs, cites)}", level=1)  # type: ignore[union-attr]
            continue
        if stripped.startswith("\\subsection"):
            flush_paragraph()
            subsection_no += 1
            title_match = re.search(r"\\subsection\{(.+?)\}", stripped)
            doc.add_heading(
                f"{section_no}.{subsection_no} {clean_tex(title_match.group(1), refs, cites)}",
                level=2,
            )  # type: ignore[union-attr]
            continue
        if stripped.startswith("\\begin{enumerate}"):
            flush_paragraph()
            list_mode = "numbered"
            continue
        if stripped.startswith("\\begin{itemize}"):
            flush_paragraph()
            list_mode = "bullet"
            continue
        if stripped.startswith("\\end{enumerate}") or stripped.startswith("\\end{itemize}"):
            flush_paragraph()
            list_mode = None
            continue
        if stripped.startswith("\\item"):
            flush_paragraph()
            add_list_item(doc, stripped.replace("\\item", "", 1), refs, cites, numbered=(list_mode == "numbered"))
            continue
        if stripped.startswith("\\") and not stripped.startswith("\\#"):
            continue
        paragraph_lines.append(stripped)

    flush_paragraph()
    add_bibliography(doc, tex, refs, cites)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    args = parse_args()
    output_path = ROOT / args.output
    build_docx(ROOT / args.tex, output_path)
    print(f"Word document written to {output_path}")


if __name__ == "__main__":
    main()
